# -*- coding: utf-8 -*-

import json
import sys

# 例如这里设置为一百万
sys.setrecursionlimit(100000000)

import docx.document
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt, RGBColor
from docx.table import _Cell
from pygments.styles.rainbow_dash import GREY


# 添加几级目录
def add_bold_heading(doc, text, level, isBold: bool, font_size=None, font_color=None):
    # 判断doc是否是Document对象
    if isinstance(doc, docx.document.Document):
        heading = doc.add_heading(text, level)
        run = heading.runs[0]
    else:  # 如果时 Paragraph
        run = doc.add_run(text)
    font = run.font
    font.bold = isBold
    if font_size:
        font.size = Pt(font_size)
    if font_color:
        font.color.rgb = RGBColor(*font_color)


def add_table_from_parameters(doc: Document, definition: dict, parameters: dict):
    # 添加表格
    headers = ["参数名称", "物理含义", "类型", "是否必填", "参数位置", "示例"]
    table = doc.add_table(rows=1, cols=len(headers))

    # 添加表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        run = header_cells[i].paragraphs[0].add_run(header)
        run.font.bold = True
        color = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(
            nsdecls('w'), color_value=GREY))
        # 设置单元格背景颜色
        header_cells[i]._tc.get_or_add_tcPr().append(color)
        # 设置单元格上下左右边框
        set_table_singleBoard(header_cells[i])

    # 添加参数数据
    for param in parameters:
        row_cells = table.add_row().cells
        row_cells[0].text = param.get('name', '')
        run = row_cells[0].paragraphs[0].runs[0]
        # Set the bold property of the run's font
        run.font.bold = True
        row_cells[1].text = param.get('description', '')
        row_cells[2].text = param.get('type', '')
        row_cells[3].text = str(param.get('required', ''))
        position = param.get('in', '')
        row_cells[4].text = position
        row_cells[5].text = str(param.get('x-example', ''))
        if position == 'body':
            originalRef = param.get('schema', {}).get('originalRef', '')
            if originalRef in definition:
                if 'description' in definition[originalRef]:
                    row_cells[1].text = definition[originalRef]['description']
                if 'title' in definition[originalRef]:
                    row_cells[2].text = definition[originalRef]['title']
                try:
                    recursive_add_parameters(table, definition, 0, originalRef, dtoName=[originalRef])
                except Exception as e:
                    print(f"Error: {e}")
                    print(originalRef)

    for row_cells in table.rows:
        for cell in row_cells.cells:
            set_table_singleBoard(cell)


def recursive_add_parameters(table, definition: dict, depth: int, obejct: str, dtoName: []):
    depth += 1
    prefix = "++" * depth
    if obejct not in definition:
        return
        # 处理对象转为参数
    requiredList = definition[obejct].get('required', [])
    for pName, pData in definition[obejct].get('properties', {}).items():
        row_cells = table.add_row().cells
        row_cells[0].text = prefix + pName
        run = row_cells[0].paragraphs[0].runs[0]
        # Set the bold property of the run's font
        run.font.bold = True
        row_cells[1].text = pData.get('description', '')
        type = pData.get('type', '')
        row_cells[2].text = type
        row_cells[3].text = 'true' if pName in requiredList else 'false'
        row_cells[5].text = 'body'
        row_cells[5].text = str(pData.get('example', ''))
        if type == 'array':
            if 'items' in pData and 'originalRef' in pData['items']:
                originalRef = pData['items']['originalRef']
                if originalRef in dtoName:
                    continue
                else:
                    dtoName.append(pName)
                row_cells[1].text = definition[originalRef].get('description', '')
                row_cells[2].text = definition[originalRef].get('title', '')
                recursive_add_parameters(table, definition, depth, originalRef, dtoName)
            else:
                row_cells[2].text = type + "(" + pData['items']['type'] + ")"
        elif type == 'obejct':
            print(pName)


def add_table_from_response_status(doc: Document, responses: dict):
    # 添加表格
    headers = ["状态码", "说明", "类型"]
    table = doc.add_table(rows=1, cols=len(headers))

    # 添加表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        run = header_cells[i].paragraphs[0].add_run(header)
        run.font.bold = True
        color = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(
            nsdecls('w'), color_value=GREY))
        # 设置单元格背景颜色
        header_cells[i]._tc.get_or_add_tcPr().append(color)
        # 设置单元格上下左右边框
        set_table_singleBoard(header_cells[i])

    # 添加参数数据
    for resp, resp_data in responses.items():
        row_cells = table.add_row().cells
        row_cells[0].text = resp
        run = row_cells[0].paragraphs[0].runs[0]
        run.font.bold = True
        row_cells[1].text = resp_data.get('description', '')

        if resp == "200":
            row_cells[2].text = "字符串|null|array|对应接口的响应数据对象"
        elif resp_data.get('schema'):
            row_cells[2].text = ("类型：" + resp_data['schema']['originalRef'])

        for cell in row_cells:
            set_table_singleBoard(cell)


def add_table_from_response_data(doc: Document, definition: dict, responses: dict):
    # 添加表格
    headers = ["参数名称", "参数说明", "类型", "示例"]
    table = doc.add_table(rows=1, cols=len(headers))

    # 添加表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        run = header_cells[i].paragraphs[0].add_run(header)
        run.font.bold = True
        color = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(
            nsdecls('w'), color_value=GREY))
        # 设置单元格背景颜色
        header_cells[i]._tc.get_or_add_tcPr().append(color)

    # 添加参数数据
    recurse_add_response_object(table, 0, definition, responses, dtoName=[])
    for row in table.rows:
        for cell in row.cells:
            # 设置单元格上下左右边框
            set_table_singleBoard(cell)


# 添加参数数据
def recurse_add_response_object(table, depth: int, definition: dict, responses: dict, dtoName: []):
    prefix = "++" * depth
    for resp, resp_data in responses.items():
        row_cells = table.add_row().cells
        row_cells[0].text = prefix + resp
        run = row_cells[0].paragraphs[0].runs[0]
        run.font.bold = True
        row_cells[1].text = resp_data.get('description', '')
        row_cells[3].text = str(resp_data.get('example', ''))
        if 'items' not in resp_data:
            row_cells[2].text = resp_data.get('type', '')
            if 'originalRef' in resp_data:
                row_cells[2].text = resp_data['originalRef']
                if resp_data['originalRef'] not in dtoName:
                    dtoName.append(resp_data['originalRef'])
                    recurse_add_response_object(table, depth + 1, definition,
                                                definition[resp_data['originalRef']]['properties'], dtoName)
            continue
        if 'originalRef' in resp_data['items']:
            row_cells[2].text = "array(" + resp_data['items']['originalRef'] + ")"
            if resp_data['items']['originalRef'] not in dtoName:
                dtoName.append(resp_data['items']['originalRef'])
                recurse_add_response_object(table, depth + 1, definition,
                                            definition[resp_data['items']['originalRef']]['properties'], dtoName)
        else:
            t = resp_data.get('type', '')
            if t == "array":
                row_cells[2].text = "array"
                if 'type' in resp_data['items'] and resp_data['items']['type'] == "string":
                    row_cells[2].text = "array(String)"
                else:
                    row_cells[2].text = "array(" + resp_data['items']['type'] + ")"
            elif t == "object":
                row_cells[2].text = "object(未知对象)"


# 设置 table 的边框，用法与 cell 类似
def set_table_boarder(table, **kwargs):
    """
    Set table`s border
    Usage:
    set_table_border(
        table,
        top={"sz": 12, "val": "single", "color": "#FF0000"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    borders = OxmlElement('w:tblBorders')
    for tag in ('bottom', 'top', 'left', 'right', 'insideV', 'insideH'):
        edge_data = kwargs.get(tag)
        if edge_data:
            any_border = OxmlElement(f'w:{tag}')
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    any_border.set(qn(f'w:{key}'), str(edge_data[key]))
            borders.append(any_border)
            table._tbl.tblPr.append(borders)


# 设置cell的边框
def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


# 将table 的所有单元格四个边设置为 0.5 镑, 黑色, 实线
def set_table_singleBoard(cell): return set_cell_border(
    cell, top={"sz": 4, "val": "single", "color": "#000000"},
    bottom={"sz": 4, "val": "single", "color": "#000000"},
    left={"sz": 4, "val": "single", "color": "#000000"},
    right={"sz": 4, "val": "single", "color": "#000000"},
    insideV={"sz": 4, "val": "single", "color": "#000000"},
    insideH={"sz": 4, "val": "single", "color": "#000000"})


# 示例用法
parameters_json = """
{
  "parameters": [
    {
      "name": "assetClassify",
      "in": "query",
      "description": "资产分类",
      "required": false,
      "type": "string"
    },
    {
      "name": "current",
      "in": "query",
      "description": "当前页面,大于等于1",
      "required": false,
      "type": "integer",
      "format": "int32",
      "x-example": 1
    }
  ]
}
"""
document = Document()
add_bold_heading(document, "参数表格", level=1, isBold=True, font_size=14, font_color=(0, 0, 255))
parameters_data = json.loads(parameters_json).get('parameters', [])
add_table_from_parameters(document, {}, parameters_data)
document.save("parameter_table.docx")
