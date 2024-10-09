#!/usr/bin/env python
# -*- coding: utf-8 -*-


import argparse
import os, re

import docx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

java_files = []


# 递归读取所有的java文件
def read_java_files(directory, java_files, file_type):
    if file_type == None:
        return
    else:
        for root, dirs, files in os.walk(directory):
            # 处理文件
            for file in files:
                fileSuffix = file.split(".")[-1]
                if fileSuffix in file_type:
                    java_file = os.path.join(root + "\\" + file)
                    java_files.append(java_file)
            # 处理目录
            for dir in dirs:
                read_java_files(dir, java_files, file_type)


def set_text_style(doc, text):
    # 添加一段文本
    paragraph = doc.add_paragraph()
    # 设置段落对齐方式（可选）
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # 设置段后间隔
    paragraph.paragraph_format.space_after = Pt(2)  # 设置为10磅
    # 设置段前间隔
    paragraph.paragraph_format.space_before = Pt(0)  # 设置为0磅
    # 设置段落的行高
    paragraph.paragraph_format.line_spacing = Pt(10)
    # 设置段落之间的间距
    # paragraph.space_after = Inches(0)  # 设置行之间的间距为 0英寸
    # 添加文本到段落中
    run = paragraph.add_run(text)
    # 设置字体大小
    run.font.size = Pt(9)
    # 设置字体颜色
    run.font.color.rgb = docx.shared.RGBColor(0x0, 0x0, 0x0)  # 黑色
    # 设置字体
    run.font.name = "monospace"


def remove_java_coment(file_contents):
    # 移除单行注释
    code = re.sub(r'//.*', '', file_contents)
    # 移除多行注释
    code = re.sub(r'/\*.*?\*/', '', code, flags=re.DOTALL)
    return code


def remove_blank_lines(file_contents):
    # 将文本按行分割，并去除每行首尾的空白字符
    lines = [line for line in file_contents.splitlines() if line.strip()]
    # 将非空行重新组合成一个字符串
    return '\n'.join(lines)


def remove_vue_comment(file_contents):
    # 移除单行注释
    code = re.sub(r'//.*', '', file_contents)
    # 移除多行注释

    # 移除 <----> 包围的多行注释
    # 注意：这个正则表达式假设 <----> 不会出现在字符串内部
    code = re.sub(r'<\-\-\-+[^>]*\-\->', '', code, flags=re.DOTALL)
    return code


# 读取每一个java文件
def write_to_word(java_files, output_path, root, isWriteDir=True, noComment=False):
    doc = Document()
    dir = []

    for file in java_files:
        with open(file, "r", encoding="utf-8") as f:
            file_contents = f.read()
        # 删除注释
        if noComment:
            if file.endswith(".java"):
                file_contents = remove_java_coment(file_contents)
            elif file.endswith(".vue"):
                file_contents = remove_vue_comment(file_contents)
            # 正则匹配去掉空白行和注释
            file_contents = remove_blank_lines(file_contents)
        if isWriteDir:
            fileName = file.replace(root + "\\", '')
            # level = fileName.count("\\")
            dirs = fileName.split("\\")
            for i, d in enumerate(dirs):
                if d not in dir:
                    dir.append(d)
                    i = i if i < 9 else 8
                    i = i + 1
                    doc.add_heading(d, level=i)

        paragraphs = file_contents.split("\n")
        fileLine = len(paragraphs)
        lastWwrite = ""
        for i in range(fileLine):
            paragraph = paragraphs[i].rstrip()
            if (i + 1) < fileLine - 1:
                nextParagraph = paragraphs[i + 1].strip()
            else:
                nextParagraph = ""
            # 处理内容，去掉 import、package、get、set方法，多个空白的换行符省略成一个换行符
            if not paragraph.startswith("import") and not paragraph.startswith("package") and not (
                    nextParagraph == paragraph and paragraph == "\n") \
                    and not (lastWwrite == paragraph and lastWwrite == ''):
                set_text_style(doc, paragraph)
                lastWwrite = paragraph

        doc.save(output_path)


# 指定目录路径和输出路径
directory_path = "D:\Java\project\jcgl\stc-jcgl\src\main\java\com\cstc\jcgl\controller"
output_file_path = "output.docx"


def parseParameter():
    # 创建参数解析器,创建解析器时指定 stdout 参数为 sys.stdout。这样可以确保打包后的可执行文件正确处理标准输出。
    parser = argparse.ArgumentParser()
    parser.add_argument("-t", "--target", help="指定目标目录，必填")
    parser.add_argument("-o", "--output", help="指定输出文件名称，默认是output.docx")
    # 开关类型参数，如果制定了 -d 则为True
    parser.add_argument("-d", "--writeDir", help="指示是否生成目录，默认写目录，false不写目录")
    parser.add_argument('-f', '--fileType', nargs='+',
                        help='指定读取文件类型，可以-f java vue 指定多个文件，默认java文件')

    parser.add_argument('-c', '--noComment', help='是否删除注释，默认不删除注释，true删除注释换行')

    # 解析命令行参数
    args = parser.parse_args()
    # 处理参数
    if not args.target:
        print("没有指定目标目录")
        exit(1)
    if not os.path.isdir(args.target):
        print(args.target + "不是一个目录")
        exit(1)
    if not args.output:
        args.output = "output.docx"
    if args.writeDir is not None:
        args.writeDir = (args.writeDir.lower()) == "true"
    else:
        args.writeDir = True
    if args.noComment is not None:
        args.noComment = (args.noComment.lower()) == "true"
    else:
        args.noComment = False
    if not args.fileType:
        args.fileType = ['java']
    return args


if __name__ == '__main__':
    args = parseParameter()
    read_java_files(args.target, java_files, args.fileType)
    # # 读取每一个文件写入到word文件中
    write_to_word(java_files, args.output, args.target, args.writeDir, args.noComment)
