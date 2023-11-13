# -*- coding: utf-8 -*-


import json
import sys
from tqdm import tqdm

# 例如这里设置为一百万
sys.setrecursionlimit(100000000)

import requests
from docx import Document  # 使用 python-docx 库来生成 Word 文档
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR
from docx.shared import Pt, RGBColor

from WriteTable import add_table_from_parameters, add_bold_heading, add_table_from_response_status, \
    add_table_from_response_data


def get_api_data(api_url):
    try:
        # 发送GET请求
        response = requests.get(api_url)

        # 检查请求是否成功 (状态码为 200)
        if response.status_code == 200:
            # 返回JSON字符串
            return response.text
        else:
            print(f"请求失败，状态码: {response.status_code}")
            return None
    except Exception as e:
        print(f"发生异常: {str(e)}")
        return None


# 设置字体 48号黑色宋体居中
def set_text(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Set alignment to center
    run = paragraph.add_run(text)
    font = run.font
    font.name = 'SimSun'  # Set font to SimSun (宋体)
    font.size = Pt(24)  # Set font size to 12-point
    font.color.rgb = RGBColor(*(0, 0, 0))  # Set font color to black


def add_vertical_line(paragraph):
    # 创建一个带颜色的阴影
    run = paragraph.add_run('  ')
    run.font.highlight_color = WD_COLOR.GREEN


def generate_api_doc(swagger_json):
    # 如果是字符串解析成json对象
    if (type(swagger_json) == str):
        swagger_json = json.loads(swagger_json)

    # 创建一个Word文档
    doc = Document()

    # 写标题
    set_text(doc, swagger_json['info']['title'])
    # 写基础信息
    doc.add_heading("1、基础信息", level=1)
    doc.add_heading("1.1 访问上下文路径：http://" + swagger_json['host'] + swagger_json['basePath'], level=2)
    # 写响应状态说明
    doc.add_heading("1.2 通用响应状态：", level=2)
    next(iter((next(iter(swagger_json['paths'].values()))).values()))
    add_table_from_response_status(doc, next(iter((next(iter(swagger_json['paths'].values()))).values()))['responses'])

    # 写响应数据结构说明
    doc.add_heading("1.3 响应数据结构说明：", level=2)
    add_bold_heading(doc.add_paragraph(),
                     "A、对应接口的响应参数描述的code=200时，data字段里面的对象结构，所有接口返回数据结构如下。" +
                     "如果与通用数据结构相同则描述的是通用数据结构。", level=1, isBold=True,
                     font_size=11, font_color=(220, 20, 60))
    add_bold_heading(doc.add_paragraph(), "B、++表示第一级嵌套对象,每加深一级多两个++", level=1, isBold=True,
                     font_size=11, font_color=(220, 20, 60))
    commonStructure = '{"code": 200, "msg": "information", "data": "object"}'
    # 将 JSON 字符串转换为 Python 字典
    data = json.loads(commonStructure)
    # 格式化输出 JSON 数据 ，indent 参数用于指定缩进空格数
    formatted_json = json.dumps(data, indent=2)
    doc.add_paragraph(formatted_json)

    doc.add_heading("2、接口信息", level=1)
    tags = []
    for t in swagger_json['tags']:
        tags.append(t['name'])
    mainCount = 0
    subCount = 0
    for path, path_data in tqdm(swagger_json['paths'].items()):
        for method, method_data in path_data.items():
            if (method_data['tags'][0] in tags):
                mainCount += 1
                subCount = 1
                doc.add_heading("2." + str(mainCount) + method_data['tags'][0], level=2)
                tags.remove(method_data['tags'][0])
            else:
                subCount += 1
            number = "2." + str(mainCount) + "." + str(subCount)
            doc.add_heading(number + method_data['summary'] +
                            ("(已过期)" if 'deprecated' in method_data and method_data['deprecated'] == 'true' else ""),
                            level=3)
            p1 = doc.add_paragraph()
            add_vertical_line(p1)
            add_bold_heading(p1, "请求方式：", level=1, isBold=True, font_size=11, font_color=(0, 0, 0))
            add_bold_heading(p1, method, level=1, isBold=False, font_size=11, font_color=(0, 0, 0))
            p2 = doc.add_paragraph()
            add_vertical_line(p2)
            add_bold_heading(p2, "请求路径：", level=1, isBold=True, font_size=11, font_color=(0, 0, 0))
            add_bold_heading(p2, path, level=1, isBold=False, font_size=11, font_color=(0, 0, 0))

            px = doc.add_paragraph()
            add_vertical_line(px)
            add_bold_heading(px, "请求数据类型：", level=1, isBold=True, font_size=11, font_color=(0, 0, 0))
            if 'consumes' in method_data:
                add_bold_heading(px, str().join(method_data['consumes']), level=1, isBold=False, font_size=11,
                                 font_color=(0, 0, 0))

            pxy = doc.add_paragraph()
            add_vertical_line(pxy)
            add_bold_heading(pxy, "响应数据类型：", level=1, isBold=True, font_size=11, font_color=(0, 0, 0))
            if 'produces' in method_data:
                add_bold_heading(pxy, str().join(method_data['produces']), level=1, isBold=False, font_size=11,
                                 font_color=(0, 0, 0))

            p3 = doc.add_paragraph()
            add_vertical_line(p3)
            add_bold_heading(p3, "请求参数：", level=1, isBold=True, font_size=11, font_color=(0, 0, 0))
            # 有参数才处理
            if 'parameters' in method_data:
                try:
                    add_table_from_parameters(doc,swagger_json['definitions'], method_data['parameters'])
                except Exception as e:
                    print(f"Error: {path}, {e}")

            p4 = doc.add_paragraph()
            add_vertical_line(p4)
            add_bold_heading(p4, "响应参数：", level=1, isBold=True, font_size=11, font_color=(0, 0, 0))
            if 'responses' in method_data:
                for code, response in method_data['responses'].items():
                    if code == '200' and 'schema' in response:
                        if 'originalRef' in response['schema']:
                            add_table_from_response_data(doc, swagger_json['definitions'],
                                                         swagger_json['definitions'][response['schema']['originalRef']][
                                                             'properties'])
                        else:
                            add_bold_heading(p4, "未知响应对象(" + response['schema']['type'] + ")", level=1,
                                             isBold=False,
                                             font_size=11, font_color=(0, 0, 0))
                    break

            # 添加分页符号
            # doc.add_page_break()
            # 只获取第一个请求的方式
            break

    # 保存文档
    doc.save(swagger_json['info']['title'] + '.docx')


api_url = "http://172.16.0.97:8603/v2/api-docs?group=%E8%AE%BE%E5%A4%87%E7%AE%A1%E7%90%86"  # 一个示例API的URL
api_ry_url = "http://localhost:8603/v2/api-docs"
json_data = get_api_data(api_ry_url)
# 函数的作用是，返回传入字符串的表达式的结果,
# 注意：eval虽然方便，但是要注意安全性，就可以利用执行系统命令，删除文件等操作。，比如用户恶意输入就会获得当前目录文件
json_data = eval(json_data, type('Dummy', (dict,), dict(__getitem__=lambda s, n: n))())
generate_api_doc(json_data)

# 3、处理复杂的请求头参数
